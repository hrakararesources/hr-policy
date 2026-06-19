#!/usr/bin/env python3
"""Akara Resources — Welfare Generator v1"""
import os, re
from docx import Document
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

WELFARE_SECTIONS = [
    {"num": 1,  "title_th": "บททั่วไปและหลักการ",
     "title_en": "General Provisions and Principles",
     "start_th": "บททั่วไปและหลักการ", "end_th": "รายการสวัสดิการและสิทธิประโยชน์",
     "start_en": "General Provisions and Principles", "end_en": "Welfare and Benefits"},
    {"num": 2,  "title_th": "การตรวจสุขภาพ",
     "title_en": "Medical Examination",
     "start_th": "การตรวจสุขภาพ", "end_th": "การประกันชีวิตกลุ่ม",
     "start_en": "Medical Examination", "end_en": "Group Life Insurance"},
    {"num": 3,  "title_th": "การประกันชีวิตกลุ่ม",
     "title_en": "Group Life Insurance",
     "start_th": "การประกันชีวิตกลุ่ม", "end_th": "การประกันอุบัติเหตุกลุ่ม",
     "start_en": "Group Life Insurance", "end_en": "Group Accident Insurance"},
    {"num": 4,  "title_th": "การประกันอุบัติเหตุกลุ่ม",
     "title_en": "Group Accident Insurance",
     "start_th": "การประกันอุบัติเหตุกลุ่ม", "end_th": "เครื่องแบบพนักงาน",
     "start_en": "Group Accident Insurance", "end_en": "Employee Uniforms"},
    {"num": 5,  "title_th": "เครื่องแบบพนักงาน",
     "title_en": "Employee Uniforms",
     "start_th": "เครื่องแบบพนักงาน", "end_th": "อาหารพนักงาน",
     "start_en": "Employee Uniforms", "end_en": "Employee Meals"},
    {"num": 6,  "title_th": "อาหารพนักงาน",
     "title_en": "Employee Meals",
     "start_th": "อาหารพนักงาน", "end_th": "ค่าเดินทางมาทำงาน",
     "start_en": "Employee Meals", "end_en": "Transportation Allowance"},
    {"num": 7,  "title_th": "ค่าเดินทางมาทำงาน",
     "title_en": "Transportation Allowance",
     "start_th": "ค่าเดินทางมาทำงาน", "end_th": "ค่ากะทำงาน",
     "start_en": "Transportation Allowance", "end_en": "Shift Allowance"},
    {"num": 8,  "title_th": "ค่ากะทำงาน",
     "title_en": "Shift Allowance",
     "start_th": "ค่ากะทำงาน", "end_th": "เงินช่วยเหลืออุปสมบท",
     "start_en": "Shift Allowance", "end_en": "Ordination Support Allowance"},
    {"num": 9,  "title_th": "เงินช่วยเหลืออุปสมบท",
     "title_en": "Ordination Support Allowance",
     "start_th": "เงินช่วยเหลืออุปสมบท", "end_th": "เงินช่วยเหลือกรณีสมรส",
     "start_en": "Ordination Support Allowance", "end_en": "Marriage Support Allowance"},
    {"num": 10, "title_th": "เงินช่วยเหลือกรณีสมรส",
     "title_en": "Marriage Support Allowance",
     "start_th": "เงินช่วยเหลือกรณีสมรส", "end_th": "เงินช่วยเหลือกรณีคลอดบุตร",
     "start_en": "Marriage Support Allowance", "end_en": "Childbirth Support Allowance"},
    {"num": 11, "title_th": "เงินช่วยเหลือกรณีคลอดบุตร",
     "title_en": "Childbirth Support Allowance",
     "start_th": "เงินช่วยเหลือกรณีคลอดบุตร", "end_th": "เงินช่วยเหลือกรณีเสียชีวิตของบุคคลในครอบครัว",
     "start_en": "Childbirth Support Allowance", "end_en": "Family Bereavement Assistance"},
    {"num": 12, "title_th": "เงินช่วยเหลือกรณีเสียชีวิตของบุคคลในครอบครัว",
     "title_en": "Family Bereavement Assistance",
     "start_th": "เงินช่วยเหลือกรณีเสียชีวิตของบุคคลในครอบครัว", "end_th": "เงินช่วยเหลือกรณีเสียชีวิตของพนักงาน",
     "start_en": "Family Bereavement Assistance", "end_en": "Death Benefit for Employees"},
    {"num": 13, "title_th": "เงินช่วยเหลือกรณีเสียชีวิตของพนักงาน",
     "title_en": "Death Benefit for Employees",
     "start_th": "เงินช่วยเหลือกรณีเสียชีวิตของพนักงาน", "end_th": "รางวัลตอบแทนการทำงานตามอายุงานต่อเนื่อง",
     "start_en": "Death Benefit for Employees", "end_en": "Long-Service Recognition Award"},
    {"num": 14, "title_th": "รางวัลตอบแทนการทำงานตามอายุงานต่อเนื่อง",
     "title_en": "Long-Service Recognition Award",
     "start_th": "รางวัลตอบแทนการทำงานตามอายุงานต่อเนื่อง", "end_th": "กองทุนสำรองเลี้ยงชีพ",
     "start_en": "Long-Service Recognition Award", "end_en": "Provident Fund"},
    {"num": 15, "title_th": "กองทุนสำรองเลี้ยงชีพ",
     "title_en": "Provident Fund",
     "start_th": "กองทุนสำรองเลี้ยงชีพ", "end_th": "สวัสดิการค่ารักษาพยาบาล",
     "start_en": "Provident Fund", "end_en": "Medical Benefits"},
    {"num": 16, "title_th": "สวัสดิการค่ารักษาพยาบาล",
     "title_en": "Medical Benefits",
     "start_th": "สวัสดิการค่ารักษาพยาบาล", "end_th": "โทรศัพท์มือถือ",
     "start_en": "Medical Benefits", "end_en": "Mobile Phone"},
    {"num": 17, "title_th": "โทรศัพท์มือถือ",
     "title_en": "Mobile Phone",
     "start_th": "โทรศัพท์มือถือ", "end_th": "ค่าเดินทางกรณีใช้รถยนต์ส่วนตัว",
     "start_en": "Mobile Phone", "end_en": "Travel Allowance for Use of Personal Vehicle"},
    {"num": 18, "title_th": "ค่าเดินทางกรณีใช้รถยนต์ส่วนตัว",
     "title_en": "Travel Allowance for Use of Personal Vehicle",
     "start_th": "ค่าเดินทางกรณีใช้รถยนต์ส่วนตัว", "end_th": "ประกาศใช้ระเบียบ",
     "start_en": "Travel Allowance for Use of Personal Vehicle", "end_en": "Announcement"},
]

SKIP_TH = ['สารบัญ','สวัสดิการและสิทธิประโยชน์พนักงาน','คู่มือ',
           'บริษัท อัครา รีซอร์สเซส','ฉบับ เดือน','รายการสวัสดิการและสิทธิประโยชน์',
           'ประกาศใช้ระเบียบ']
SKIP_EN = ['Table of Contents','Employee Welfare and Benefits','Akara Resources Public',
           'Edition:','Welfare and Benefits','Announcement']

SIZE_H1_TH = 228600
SIZE_H1_EN = 139700

def extract_section(paragraphs, start_key, end_key):
    collecting = False
    result = []
    for para in paragraphs:
        t = para.text.strip()
        if not t: continue
        if not collecting:
            if t == start_key or t.startswith(start_key):
                collecting = True
                continue
        if collecting:
            if end_key and (t == end_key or t.startswith(end_key)):
                break
            result.append(para)
    return result

def build_content(paragraphs, start_key, end_key, skip_list, size_h1):
    paras = extract_section(paragraphs, start_key, end_key)
    counters = [0] * 10
    html_parts = []

    for para in paras:
        t = para.text.strip()
        if not t: continue
        if any(s in t for s in skip_list): continue

        is_bold = any(run.bold for run in para.runs if run.text.strip())
        size = None
        for run in para.runs:
            if run.text.strip() and run.font.size:
                size = run.font.size
                break

        style = para.style.name
        numPr = para._element.find('.//' + qn('w:numPr'))

        # หัวข้อใหญ่
        if size and size >= size_h1 and is_bold:
            counters = [0] * 10
            html_parts.append(f'<h2>{t}</h2>')
            continue

        # หัวข้อ bold Normal
        if is_bold and style == 'Normal':
            html_parts.append(f'<h3>{t}</h3>')
            continue

        # List มี numbering
        if numPr is not None:
            ilvl_el = numPr.find(qn('w:ilvl'))
            level = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0
            while len(counters) <= level:
                counters.append(0)
            counters[level] += 1
            for l in range(level + 1, len(counters)):
                counters[l] = 0

            if level == 0:
                num_str = f'{counters[0]}.'
            elif level == 1:
                num_str = f'{counters[0]}.{counters[1]}'
            else:
                num_str = f'{counters[0]}.{counters[1]}.{counters[2]}'

            padding = (level + 1) * 1.5
            if is_bold:
                html_parts.append(
                    f'<p style="padding-left:{padding}em;font-weight:600">'
                    f'<span style="color:var(--blue);margin-right:6px">{num_str}</span>{t}</p>'
                )
            else:
                html_parts.append(
                    f'<p style="padding-left:{padding}em">'
                    f'<span style="color:var(--text-muted);margin-right:6px">{num_str}</span>{t}</p>'
                )
            continue

        # List ไม่มี numbering
        if style == 'List Paragraph':
            if is_bold:
                html_parts.append(f'<h4>{t}</h4>')
            else:
                html_parts.append(f'<p style="padding-left:1.5em">{t}</p>')
            continue

        html_parts.append(f'<p>{t}</p>')

    return "\n".join(html_parts) or '<p style="color:var(--text-muted)">ไม่มีเนื้อหา</p>'

def build_sidebar(sections, current_file):
    links = []
    for s in sections:
        sn = s["num"]
        sf = f"welfare-{sn:02d}.html"
        active = ' active' if sf == current_file else ''
        links.append(
            f'<a href="{sf}" class="sidebar-chapter{active}">'
            f'<span class="sidebar-num">{sn:02d}</span>'
            f'<span class="sidebar-label">'
            f'<span class="text-th">{s["title_th"]}</span>'
            f'<span class="text-en">{s["title_en"]}</span>'
            f'</span></a>'
        )
    return "\n          ".join(links)

PAGE_HTML = '''<!DOCTYPE html>
<html lang="th" data-lang="th">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title_th} — Akara Resources</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;500;600;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../assets/style.css">
  <style>
    body {{ display:flex; flex-direction:column; min-height:100vh; }}
    .page-wrap {{ display:flex; flex:1; }}
    .sidebar-fixed {{
      width:270px; flex-shrink:0; position:sticky; top:60px;
      height:calc(100vh - 60px); overflow-y:auto;
      background:#fff; border-right:1.5px solid var(--gray-2); padding:12px 8px;
    }}
    .sidebar-fixed::-webkit-scrollbar {{ width:4px; }}
    .sidebar-fixed::-webkit-scrollbar-thumb {{ background:var(--gray-3); border-radius:2px; }}
    .sidebar-home {{
      display:flex; align-items:center; gap:6px; padding:7px 10px;
      font-size:0.8rem; color:var(--text-muted); text-decoration:none;
      border-radius:6px; margin-bottom:2px;
    }}
    .sidebar-home:hover {{ background:var(--gray-1); color:var(--blue); }}
    .sidebar-section {{
      font-size:0.68rem; font-weight:700; letter-spacing:0.1em;
      text-transform:uppercase; color:var(--text-muted);
      padding:10px 10px 4px; display:block;
    }}
    .sidebar-chapter {{
      display:flex; align-items:flex-start; gap:8px; padding:8px 10px;
      border-radius:8px; text-decoration:none; color:var(--text-muted);
      font-size:0.8rem; line-height:1.35; transition:all 0.15s; margin-bottom:1px;
    }}
    .sidebar-chapter:hover {{ background:#FEF9E0; color:#7A6200; }}
    .sidebar-chapter.active {{ background:var(--gold); color:var(--dark); }}
    .sidebar-chapter.active .sidebar-num {{ background:rgba(0,0,0,0.1); color:var(--dark); }}
    .sidebar-num {{
      min-width:26px; height:26px; border-radius:6px;
      background:#FEF9E0; color:#7A6200;
      font-size:0.7rem; font-weight:700;
      display:flex; align-items:center; justify-content:center; flex-shrink:0; margin-top:1px;
    }}
    .sidebar-label {{ flex:1; }}
    .main-content {{ flex:1; padding:32px 32px 60px; min-width:0; }}
    .chapter-card {{ max-width:740px; margin:0 auto; }}
    .chapter-body h2 {{ color:#7A6200; font-size:1.1rem; margin:28px 0 8px; padding-bottom:6px; border-bottom:2px solid #FEF9E0; }}
    .chapter-body h3 {{ color:var(--dark); font-size:1rem; margin:20px 0 6px; font-weight:700; }}
    .chapter-body h4 {{ color:#7A6200; font-size:0.95rem; margin:12px 0 4px; font-weight:600; }}
    .chapter-body p {{ color:var(--text); font-size:0.95rem; margin-bottom:0.8em; line-height:1.8; }}
    @media(max-width:768px) {{
      .sidebar-fixed {{ display:none; }}
      .sidebar-fixed.open {{ display:block; position:fixed; top:60px; left:0; z-index:200; height:calc(100vh - 60px); box-shadow:4px 0 20px rgba(0,0,0,0.15); }}
      .mob-menu {{ display:flex !important; }}
    }}
  </style>
</head>
<body>
  <header class="topbar" style="background:linear-gradient(135deg,#7A6200,#B8940A,var(--gold-dark))">
    <div class="topbar-left">
      <button class="mob-menu" onclick="toggleSidebar()" style="display:none;background:rgba(255,255,255,0.15);border:none;color:#fff;padding:6px 10px;border-radius:6px;cursor:pointer;font-size:1rem;margin-right:8px;">☰</button>
      <img src="../assets/Akara Logo.jpg" alt="Akara Resources" class="topbar-logo">
      <div class="topbar-title">
        <span class="text-th">คู่มือพนักงาน</span><span class="text-en">Employee Handbook</span>
        <span>Akara Resources Public Company Limited</span>
      </div>
    </div>
    <div class="topbar-right">
      <div class="lang-toggle">
        <button class="lang-btn active" onclick="setLang(\'th\')">TH</button>
        <button class="lang-btn" onclick="setLang(\'en\')">EN</button>
      </div>
    </div>
  </header>

  <div class="page-wrap">
    <aside class="sidebar-fixed" id="sidebar">
      <a href="../index.html" class="sidebar-home">← <span class="text-th">หน้าหลัก</span><span class="text-en">Home</span></a>
      <span class="sidebar-section text-th">สวัสดิการและสิทธิประโยชน์</span>
      <span class="sidebar-section text-en">Welfare & Benefits</span>
      <nav style="display:flex;flex-direction:column;gap:1px;">
          {sidebar_links}
      </nav>
    </aside>

    <main class="main-content">
      <article class="chapter-card">
        <div class="chapter-badge" style="background:#FEF9E0;color:#7A6200">
          <span class="text-th">หัวข้อที่ {num}</span>
          <span class="text-en">Section {num}</span>
        </div>
        <h1 class="chapter-title">
          <span class="text-th">{title_th}</span>
          <span class="text-en">{title_en}</span>
        </h1>
        <div class="chapter-divider" style="background:linear-gradient(90deg,#7A6200,var(--gold))"></div>
        <div class="chapter-body">
          <div class="text-th">{content_th}</div>
          <div class="text-en">{content_en}</div>
        </div>
        <nav class="chapter-nav">{prev_link}{next_link}</nav>
      </article>
    </main>
  </div>

  <footer class="footer">
    <span class="text-th">© 2026 บริษัท อัครา รีซอร์สเซส จำกัด (มหาชน) — ฝ่ายทรัพยากรบุคคล</span>
    <span class="text-en">© 2026 Akara Resources Public Company Limited — Human Resources Department</span>
  </footer>

  <script>
    function setLang(lang) {{
      document.documentElement.setAttribute(\'data-lang\', lang);
      document.querySelectorAll(\'.lang-btn\').forEach(b => b.classList.remove(\'active\'));
      document.querySelector(`.lang-btn[onclick="setLang(\'${{lang}}\')"]`).classList.add(\'active\');
      localStorage.setItem(\'hr-lang\', lang);
    }}
    const saved = localStorage.getItem(\'hr-lang\');
    if (saved) setLang(saved);
    function toggleSidebar() {{
      document.getElementById(\'sidebar\').classList.toggle(\'open\');
    }}
  </script>
</body>
</html>'''

def main():
    print("="*56)
    print("  Akara Resources — Welfare Generator v1")
    print("="*56)

    docx_th = os.path.join(BASE_DIR, "content", "welfare", "01_Welfare and Benefits Regulation TH.docx")
    docx_en = os.path.join(BASE_DIR, "content", "welfare", "02_Welfare and Benefits Regulation EN.docx")
    out_dir = os.path.join(BASE_DIR, "welfare")

    if not os.path.exists(docx_th):
        print(f"\n❌ ไม่พบ: {docx_th}"); return

    print(f"\n📄 อ่าน TH docx...")
    paras_th = Document(docx_th).paragraphs
    print(f"   {len(paras_th)} paragraphs")

    paras_en = None
    if os.path.exists(docx_en):
        print(f"📄 อ่าน EN docx...")
        paras_en = Document(docx_en).paragraphs
        print(f"   {len(paras_en)} paragraphs")

    print(f"\n✍️  สร้าง HTML {len(WELFARE_SECTIONS)} หน้า...")
    total = len(WELFARE_SECTIONS)

    for i, sec in enumerate(WELFARE_SECTIONS):
        sn = sec["num"]
        filename = f"welfare-{sn:02d}.html"

        content_th = build_content(paras_th, sec["start_th"], sec["end_th"], SKIP_TH, SIZE_H1_TH)
        content_en = build_content(paras_en, sec["start_en"], sec["end_en"], SKIP_EN, SIZE_H1_EN) if paras_en else content_th

        sidebar_links = build_sidebar(WELFARE_SECTIONS, filename)

        prev_link = next_link = ""
        if i > 0:
            ps = WELFARE_SECTIONS[i-1]
            prev_link = f'<a href="welfare-{ps["num"]:02d}.html">← <span class="text-th">{ps["title_th"][:30]}</span><span class="text-en">{ps["title_en"][:30]}</span></a>'
        if i < total - 1:
            ns = WELFARE_SECTIONS[i+1]
            next_link = f'<a href="welfare-{ns["num"]:02d}.html" class="next"><span class="text-th">{ns["title_th"][:30]}</span><span class="text-en">{ns["title_en"][:30]}</span> →</a>'

        html = PAGE_HTML.format(
            title_th=sec["title_th"], title_en=sec["title_en"],
            num=sn, content_th=content_th, content_en=content_en,
            prev_link=prev_link, next_link=next_link,
            sidebar_links=sidebar_links
        )

        with open(os.path.join(out_dir, filename), "w", encoding="utf-8") as f:
            f.write(html)
        print(f"  ✅ {filename}")

    print(f"\n✅ เสร็จ! {total} หน้า")
    print("\n" + "="*56)
    print("  git add . && git commit -m 'add welfare' && git push")
    print("="*56 + "\n")

if __name__ == "__main__":
    main()
