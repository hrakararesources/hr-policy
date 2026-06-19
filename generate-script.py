#!/usr/bin/env python3
"""Akara Resources — HR Policy Generator v7 (TH+EN docx)"""
import os, re
from docx import Document
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

WORK_RULES_CHAPTERS = [
    {"num": 0,  "title_th": "สารจากผู้บริหาร",       "title_en": "Message from Management",
     "start_th": "สารจากผู้บริหาร", "end_th": "หมวดที่ 1",
     "start_en": "Message from the Management", "end_en": "Chapter 1"},
    {"num": 1,  "title_th": "บททั่วไป",               "title_en": "General Provisions",
     "start_th": "หมวดที่ 1",  "end_th": "หมวดที่ 2",
     "start_en": "Chapter 1",  "end_en": "Chapter 2"},
    {"num": 2,  "title_th": "วิสัยทัศน์องค์กร เสาหลักกลยุทธ์และนโยบายองค์กร",
     "title_en": "Corporate Vision, Strategic Pillars and Policies",
     "start_th": "หมวดที่ 2",  "end_th": "หมวดที่ 3",
     "start_en": "Chapter 2",  "end_en": "Chapter 3"},
    {"num": 3,  "title_th": "การว่าจ้าง",  "title_en": "Employment",
     "start_th": "หมวดที่ 3",  "end_th": "หมวดที่ 4",
     "start_en": "Chapter 3",  "end_en": "Chapter 4"},
    {"num": 4,  "title_th": "การสับเปลี่ยนหน้าที่ การโยกย้าย และการปฏิบัติงานที่บ้าน",
     "title_en": "Job Rotation, Transfer and Work from Home",
     "start_th": "หมวดที่ 4",  "end_th": "หมวดที่ 5",
     "start_en": "Chapter 4",  "end_en": "Chapter 5"},
    {"num": 5,  "title_th": "วันทำงาน เวลาทำงานปกติ และเวลาพัก",
     "title_en": "Working Days, Hours and Rest Periods",
     "start_th": "หมวดที่ 5",  "end_th": "หมวดที่ 6",
     "start_en": "Chapter 5",  "end_en": "Chapter 6"},
    {"num": 6,  "title_th": "วันหยุด และหลักเกณฑ์วันหยุด",
     "title_en": "Holidays and Holiday Criteria",
     "start_th": "หมวดที่ 6",  "end_th": "หมวดที่ 7",
     "start_en": "Chapter 6",  "end_en": "Chapter 7"},
    {"num": 7,  "title_th": "หลักเกณฑ์การทำงานล่วงเวลา และการทำงานในวันหยุด",
     "title_en": "Overtime and Holiday Work Criteria",
     "start_th": "หมวดที่ 7",  "end_th": "หมวดที่ 8",
     "start_en": "Chapter 7",  "end_en": "Chapter 8"},
    {"num": 8,  "title_th": "การจ่ายค่าจ้าง ค่าล่วงเวลา และค่าทำงานในวันหยุด",
     "title_en": "Wages, Overtime Pay and Holiday Pay",
     "start_th": "หมวดที่ 8",  "end_th": "หมวดที่ 9",
     "start_en": "Chapter 8",  "end_en": "Chapter 9"},
    {"num": 9,  "title_th": "วันลา และหลักเกณฑ์การลา",
     "title_en": "Leave and Leave Criteria",
     "start_th": "หมวดที่ 9",  "end_th": "หมวดที่ 10",
     "start_en": "Chapter 9",  "end_en": "Chapter 10"},
    {"num": 10, "title_th": "วินัย และโทษทางวินัย",
     "title_en": "Discipline and Disciplinary Penalties",
     "start_th": "หมวดที่ 10", "end_th": "หมวดที่ 11",
     "start_en": "Chapter 10", "end_en": "Chapter 11"},
    {"num": 11, "title_th": "การร้องทุกข์",
     "title_en": "Grievance Procedure",
     "start_th": "หมวดที่ 11", "end_th": "หมวดที่ 12",
     "start_en": "Chapter 11", "end_en": "Chapter 12"},
    {"num": 12, "title_th": "การพ้นสภาพการเป็นพนักงาน",
     "title_en": "Termination of Employment",
     "start_th": "หมวดที่ 12", "end_th": "หมวดที่ 13",
     "start_en": "Chapter 12", "end_en": "Chapter 13"},
    {"num": 13, "title_th": "สวัสดิการและสิทธิประโยชน์อื่น ๆ",
     "title_en": "Welfare and Other Benefits",
     "start_th": "หมวดที่ 13", "end_th": "หมวดที่ 14",
     "start_en": "Chapter 13", "end_en": "Chapter 14"},
    {"num": 14, "title_th": "ข้อกำหนดทั่วไปและการบังคับใช้",
     "title_en": "General Provisions and Enforcement",
     "start_th": "หมวดที่ 14", "end_th": None,
     "start_en": "Chapter 14", "end_en": None},
]

SKIP_TH = ['สารบัญ','ระเบียบข้อบังคับเกี่ยวกับการทำงาน','บริษัท อัครา รีซอร์สเซส','ฉบับ เดือน','ประเภทกิจการ','ทำเหมืองแร่']
SKIP_EN = ['Table of Contents','Work Rules and Regulations','Akara Resources Public Company','March 2026','Type of Business','Mining Operations']

SIZE_H1_TH = 228600
SIZE_H1_EN = 152400

def para_to_html(para, skip_list, size_h1):
    t = para.text.strip()
    if not t: return None
    if any(s in t for s in skip_list): return None

    is_bold = any(run.bold for run in para.runs if run.text.strip())
    size = None
    for run in para.runs:
        if run.text.strip() and run.font.size:
            size = run.font.size
            break

    style = para.style.name

    # หัวข้อหมวด — ตัวใหญ่ bold
    if size and size >= size_h1 and is_bold:
        return f'<h2>{t}</h2>'

    # หัวข้อ bold Normal style
    if is_bold and style == 'Normal':
        return f'<h3>{t}</h3>'

    # List Paragraph — ดู numPr ว่ามีเลขกำกับไหม
    if style == 'List Paragraph':
        # ตรวจว่ามี numbering หรือเปล่า
        numPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        if numPr is not None:
            # มีเลขกำกับ — ใช้ <ol> style
            indent = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            level = int(indent.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 0)) if indent is not None else 0
            padding = (level + 1) * 1.5
            if is_bold:
                return f'<p style="padding-left:{padding}em;font-weight:600">{t}</p>'
            return f'<p style="padding-left:{padding}em">{t}</p>'
        else:
            # ไม่มีเลข ไม่มี bullet — paragraph ธรรมดา indent
            if is_bold:
                return f'<p style="padding-left:1.5em;font-weight:600">{t}</p>'
            return f'<p style="padding-left:1.5em">{t}</p>'

    # Normal paragraph
    return f'<p>{t}</p>'

def wrap_lists(html_parts):
    result = []
    in_list = False
    for item in html_parts:
        if item and item.startswith('<li>'):
            if not in_list:
                result.append('<ul>')
                in_list = True
            result.append(item)
        else:
            if in_list:
                result.append('</ul>')
                in_list = False
            if item:
                result.append(item)
    if in_list:
        result.append('</ul>')
    return result

def extract_chapter(paragraphs, start_key, end_key, skip_list, size_h1, include_start=False):
    collecting = False
    result = []
    for para in paragraphs:
        t = para.text.strip()
        if not t: continue
        if not collecting:
            if t == start_key or t.startswith(start_key):
                collecting = True
                if include_start:
                    result.append(para)
                continue
        if collecting:
            if end_key and (t == end_key or t.startswith(end_key)):
                break
            result.append(para)
    return result

def build_content(paragraphs, start_key, end_key, skip_list, size_h1):
    paras = extract_chapter(paragraphs, start_key, end_key, skip_list, size_h1)
    
    # counters สำหรับ auto numbering
    counters = [0, 0, 0, 0, 0]
    html_parts = []
    
    for para in paras:
        t = para.text.strip()
        if not t: continue
        if any(s in t for s in skip_list): continue

        is_bold = any(run.bold for run in para.runs if run.text.strip())
        size = None
        for run in para.runs:
            if run.text.strip() and run.font.size:
                size = run.font.size
                break

        style = para.style.name
        numPr = para._element.find('.//' + qn('w:numPr'))

        # หัวข้อหมวด — reset counters
        if size and size >= size_h1 and is_bold:
            counters = [0, 0, 0]
            html_parts.append(f'<h2>{t}</h2>')
            continue

        # หัวข้อ bold Normal
        if is_bold and style == 'Normal':
            html_parts.append(f'<h3>{t}</h3>')
            continue

        # List Paragraph มี numbering
        if numPr is not None:
            ilvl_el = numPr.find(qn('w:ilvl'))
            level = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0
            
            # เพิ่ม counter ตาม level และ reset level ที่ต่ำกว่า
            counters[level] += 1
            for l in range(level + 1, 5):
                counters[l] = 0

            # สร้างเลข
            if level == 0:
                num_str = f'{counters[0]}.'
            elif level == 1:
                num_str = f'{counters[0]}.{counters[1]}'
            else:
                num_str = f'{counters[0]}.{counters[1]}.{counters[2]}'

            padding = (level + 1) * 1.5
            if is_bold:
                html_parts.append(
                    f'<p style="padding-left:{padding}em;font-weight:600">'
                    f'<span style="color:var(--blue);margin-right:6px">{num_str}</span>{t}</p>'
                )
            else:
                html_parts.append(
                    f'<p style="padding-left:{padding}em">'
                    f'<span style="color:var(--text-muted);margin-right:6px">{num_str}</span>{t}</p>'
                )
            continue

        # Normal paragraph
        html_parts.append(f'<p>{t}</p>')

    return "\n".join(html_parts) or '<p style="color:var(--text-muted)">ไม่มีเนื้อหา</p>'

def build_sidebar(chapters, current_file):
    links = []
    for c in chapters:
        cn = c["num"]
        cf = "chapter-00.html" if cn == 0 else f"chapter-{cn:02d}.html"
        num_label = "00" if cn == 0 else f"{cn:02d}"
        active = ' active' if cf == current_file else ''
        links.append(
            f'<a href="{cf}" class="sidebar-chapter{active}">'
            f'<span class="sidebar-num">{num_label}</span>'
            f'<span class="sidebar-label">'
            f'<span class="text-th">{c["title_th"]}</span>'
            f'<span class="text-en">{c["title_en"]}</span>'
            f'</span></a>'
        )
    return "\n          ".join(links)

PAGE_HTML = '''<!DOCTYPE html>
<html lang="th" data-lang="th">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title_th} — Akara Resources</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;500;600;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../assets/style.css">
  <style>
    body {{ display:flex; flex-direction:column; min-height:100vh; }}
    .page-wrap {{ display:flex; flex:1; }}
    .sidebar-fixed {{
      width:270px; flex-shrink:0; position:sticky; top:60px;
      height:calc(100vh - 60px); overflow-y:auto;
      background:#fff; border-right:1.5px solid var(--gray-2); padding:12px 8px;
    }}
    .sidebar-fixed::-webkit-scrollbar {{ width:4px; }}
    .sidebar-fixed::-webkit-scrollbar-thumb {{ background:var(--gray-3); border-radius:2px; }}
    .sidebar-home {{
      display:flex; align-items:center; gap:6px; padding:7px 10px;
      font-size:0.8rem; color:var(--text-muted); text-decoration:none;
      border-radius:6px; margin-bottom:2px;
    }}
    .sidebar-home:hover {{ background:var(--gray-1); color:var(--blue); }}
    .sidebar-section {{
      font-size:0.68rem; font-weight:700; letter-spacing:0.1em;
      text-transform:uppercase; color:var(--text-muted);
      padding:10px 10px 4px; display:block;
    }}
    .sidebar-chapter {{
      display:flex; align-items:flex-start; gap:8px; padding:8px 10px;
      border-radius:8px; text-decoration:none; color:var(--text-muted);
      font-size:0.8rem; line-height:1.35; transition:all 0.15s; margin-bottom:1px;
    }}
    .sidebar-chapter:hover {{ background:var(--blue-light); color:var(--blue); }}
    .sidebar-chapter.active {{ background:var(--blue); color:#fff; }}
    .sidebar-chapter.active .sidebar-num {{ background:rgba(255,255,255,0.25); color:#fff; }}
    .sidebar-num {{
      min-width:26px; height:26px; border-radius:6px;
      background:var(--blue-light); color:var(--blue);
      font-size:0.7rem; font-weight:700;
      display:flex; align-items:center; justify-content:center; flex-shrink:0; margin-top:1px;
    }}
    .sidebar-label {{ flex:1; }}
    .main-content {{ flex:1; padding:32px 32px 60px; min-width:0; }}
    .chapter-card {{ max-width:740px; margin:0 auto; }}
    .chapter-body h2 {{ color:var(--blue); font-size:1.1rem; margin:28px 0 8px; padding-bottom:6px; border-bottom:2px solid var(--blue-light); }}
    .chapter-body h3 {{ color:var(--dark); font-size:1rem; margin:20px 0 6px; font-weight:700; }}
    .chapter-body h4 {{ color:var(--dark); font-size:0.95rem; margin:12px 0 4px; font-weight:600; }}
    .chapter-body p {{ color:var(--text); font-size:0.95rem; margin-bottom:0.8em; line-height:1.8; }}
    .chapter-body ul {{ padding-left:1.4em; margin-bottom:1em; }}
    .chapter-body li {{ font-size:0.95rem; margin-bottom:0.5em; color:var(--text); line-height:1.7; }}
    @media(max-width:768px) {{
      .sidebar-fixed {{ display:none; }}
      .sidebar-fixed.open {{ display:block; position:fixed; top:60px; left:0; z-index:200; height:calc(100vh - 60px); box-shadow:4px 0 20px rgba(0,0,0,0.15); }}
      .mob-menu {{ display:flex !important; }}
    }}
  </style>
</head>
<body>
  <header class="topbar">
    <div class="topbar-left">
      <button class="mob-menu" onclick="toggleSidebar()" style="display:none;background:rgba(255,255,255,0.15);border:none;color:#fff;padding:6px 10px;border-radius:6px;cursor:pointer;font-size:1rem;margin-right:8px;">☰</button>
      <img src="../assets/Akara Logo.jpg" alt="Akara Resources" class="topbar-logo">
      <div class="topbar-title">
        <span class="text-th">คู่มือพนักงาน</span><span class="text-en">Employee Handbook</span>
        <span>Akara Resources Public Company Limited</span>
      </div>
    </div>
    <div class="topbar-right">
      <div class="lang-toggle">
        <button class="lang-btn active" onclick="setLang(\'th\')">TH</button>
        <button class="lang-btn" onclick="setLang(\'en\')">EN</button>
      </div>
    </div>
  </header>

  <div class="page-wrap">
    <aside class="sidebar-fixed" id="sidebar">
      <a href="../index.html" class="sidebar-home">← <span class="text-th">หน้าหลัก</span><span class="text-en">Home</span></a>
      <a href="index.html" class="sidebar-home" style="color:var(--blue);font-weight:600;">☰ <span class="text-th">สารบัญ</span><span class="text-en">Contents</span></a>
      <span class="sidebar-section text-th">ระเบียบข้อบังคับการทำงาน</span>
      <span class="sidebar-section text-en">Work Rules & Regulations</span>
      <nav style="display:flex;flex-direction:column;gap:1px;">
          {sidebar_links}
      </nav>
    </aside>

    <main class="main-content">
      <article class="chapter-card">
        <div class="chapter-badge">
          <span class="text-th">{badge_th}</span>
          <span class="text-en">{badge_en}</span>
        </div>
        <h1 class="chapter-title">
          <span class="text-th">{title_th}</span>
          <span class="text-en">{title_en}</span>
        </h1>
        <div class="chapter-divider"></div>
        <div class="chapter-body">
          <div class="text-th">{content_th}</div>
          <div class="text-en">{content_en}</div>
        </div>
        <nav class="chapter-nav">{prev_link}{next_link}</nav>
      </article>
    </main>
  </div>

  <footer class="footer">
    <span class="text-th">© 2026 บริษัท อัครา รีซอร์สเซส จำกัด (มหาชน) — ฝ่ายทรัพยากรบุคคล</span>
    <span class="text-en">© 2026 Akara Resources Public Company Limited — Human Resources Department</span>
  </footer>

  <script>
    function setLang(lang) {{
      document.documentElement.setAttribute(\'data-lang\', lang);
      document.querySelectorAll(\'.lang-btn\').forEach(b => b.classList.remove(\'active\'));
      document.querySelector(`.lang-btn[onclick="setLang(\'${{lang}}\')"]`).classList.add(\'active\');
      localStorage.setItem(\'hr-lang\', lang);
    }}
    const saved = localStorage.getItem(\'hr-lang\');
    if (saved) setLang(saved);
    function toggleSidebar() {{
      document.getElementById(\'sidebar\').classList.toggle(\'open\');
    }}
  </script>
</body>
</html>'''

TOC_ITEM = '''<a href="{filename}" class="toc-item">
  <div class="toc-num"><span class="text-th">{badge_th}</span><span class="text-en">{badge_en}</span></div>
  <div class="toc-text">
    <strong class="text-th">{title_th}</strong>
    <strong class="text-en">{title_en}</strong>
  </div>
  <span class="toc-chevron">›</span>
</a>'''

def main():
    print("="*56)
    print("  Akara Resources — HR Policy Generator v7 (TH+EN)")
    print("="*56)

    docx_th = os.path.join(BASE_DIR, "content", "work-rules", "01_Work Rules and Regulations TH.docx")
    docx_en = os.path.join(BASE_DIR, "content", "work-rules", "02_Work Rules and Regulations EN.docx")
    out_dir = os.path.join(BASE_DIR, "work-rules")
    toc_path = os.path.join(out_dir, "index.html")

    if not os.path.exists(docx_th):
        print(f"\n❌ ไม่พบ: {docx_th}"); return

    print(f"\n📄 อ่าน TH docx...")
    paras_th = Document(docx_th).paragraphs
    print(f"   {len(paras_th)} paragraphs")

    paras_en = None
    if os.path.exists(docx_en):
        print(f"📄 อ่าน EN docx...")
        paras_en = Document(docx_en).paragraphs
        print(f"   {len(paras_en)} paragraphs")
    else:
        print(f"⚠️  ไม่พบ EN docx — จะใช้ TH แทน")

    print(f"\n✍️  สร้าง HTML {len(WORK_RULES_CHAPTERS)} หน้า...")
    toc_items = []
    total = len(WORK_RULES_CHAPTERS)

    for i, ch in enumerate(WORK_RULES_CHAPTERS):
        chnum = ch["num"]
        badge_th = "สารจากผู้บริหาร" if chnum == 0 else f"หมวดที่ {chnum}"
        badge_en = "Management Message" if chnum == 0 else f"Chapter {chnum}"
        filename = "chapter-00.html" if chnum == 0 else f"chapter-{chnum:02d}.html"

        content_th = build_content(paras_th, ch["start_th"], ch["end_th"], SKIP_TH, SIZE_H1_TH)
        if paras_en:
            content_en = build_content(paras_en, ch["start_en"], ch["end_en"], SKIP_EN, SIZE_H1_EN)
        else:
            content_en = content_th

        sidebar_links = build_sidebar(WORK_RULES_CHAPTERS, filename)

        prev_link = next_link = ""
        if i > 0:
            pch = WORK_RULES_CHAPTERS[i-1]
            pn = pch["num"]
            pf = "chapter-00.html" if pn == 0 else f"chapter-{pn:02d}.html"
            prev_link = f'<a href="{pf}">← <span class="text-th">{pch["title_th"][:30]}</span><span class="text-en">{pch["title_en"][:30]}</span></a>'
        if i < total - 1:
            nch = WORK_RULES_CHAPTERS[i+1]
            nn = nch["num"]
            nf = "chapter-00.html" if nn == 0 else f"chapter-{nn:02d}.html"
            next_link = f'<a href="{nf}" class="next"><span class="text-th">{nch["title_th"][:30]}</span><span class="text-en">{nch["title_en"][:30]}</span> →</a>'

        html = PAGE_HTML.format(
            title_th=ch["title_th"], title_en=ch["title_en"],
            badge_th=badge_th, badge_en=badge_en,
            content_th=content_th, content_en=content_en,
            prev_link=prev_link, next_link=next_link,
            sidebar_links=sidebar_links
        )

        with open(os.path.join(out_dir, filename), "w", encoding="utf-8") as f:
            f.write(html)
        print(f"  ✅ {filename}")

        toc_items.append(TOC_ITEM.format(
            filename=filename, badge_th=badge_th, badge_en=badge_en,
            title_th=ch["title_th"], title_en=ch["title_en"]
        ))

    # Reset และ inject TOC ใหม่
    # เขียน index.html ใหม่ทั้งไฟล์ทุกครั้ง
    toc_content = f'''<!DOCTYPE html>
<html lang="th" data-lang="th">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>ระเบียบข้อบังคับการทำงาน — Akara Resources</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=Sarabun:wght@400;500;600;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../assets/style.css">
</head>
<body>
  <header class="topbar">
    <div class="topbar-left">
      <a href="../index.html" class="back-btn">
        <svg viewBox="0 0 14 14" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round"><path d="M9 2L4 7L9 12"/></svg>
        <span class="text-th">หน้าหลัก</span><span class="text-en">Home</span>
      </a>
      <img src="../assets/Akara Logo.jpg" alt="Akara Resources" class="topbar-logo">
      <div class="topbar-title">
        <span class="text-th">คู่มือพนักงาน</span><span class="text-en">Employee Handbook</span>
        <span>Akara Resources Public Company Limited</span>
      </div>
    </div>
    <div class="topbar-right">
      <div class="lang-toggle">
        <button class="lang-btn active" onclick="setLang('th')">TH</button>
        <button class="lang-btn" onclick="setLang('en')">EN</button>
      </div>
    </div>
  </header>
  <section class="hero" style="padding:36px 24px 32px">
    <div class="hero-eyebrow">Work Rules & Regulations</div>
    <h1>
      <span class="text-th">ระเบียบข้อบังคับการทำงาน</span>
      <span class="text-en">Work Rules & Regulations</span>
    </h1>
    <p>
      <span class="text-th">บริษัท อัครา รีซอร์สเซส จำกัด (มหาชน) — 14 หมวด</span>
      <span class="text-en">Akara Resources Public Company Limited — 14 Chapters</span>
    </p>
  </section>
  <main class="container" style="max-width:720px">
    <div class="toc-search-wrap">
      <span class="toc-search-icon">🔍</span>
      <input type="search" class="toc-search" id="tocSearch"
        placeholder="ค้นหาหัวข้อ... / Search..."
        oninput="filterTOC(this.value)">
    </div>
    <div class="toc-list" id="tocList">
{toc_block}
    </div>
    <p id="noResult" style="display:none;text-align:center;color:var(--text-muted);padding:40px 0">
      <span class="text-th">ไม่พบหัวข้อที่ค้นหา</span>
      <span class="text-en">No results found</span>
    </p>
  </main>
  <footer class="footer">
    <span class="text-th">© 2026 บริษัท อัครา รีซอร์สเซส จำกัด (มหาชน) — ฝ่ายทรัพยากรบุคคล</span>
    <span class="text-en">© 2026 Akara Resources Public Company Limited — Human Resources Department</span>
  </footer>
  <script>
    function setLang(lang) {{
      document.documentElement.setAttribute('data-lang', lang);
      document.querySelectorAll('.lang-btn').forEach(b => b.classList.remove('active'));
      document.querySelector(`.lang-btn[onclick="setLang('${{lang}}')"]`).classList.add('active');
      localStorage.setItem('hr-lang', lang);
    }}
    const saved = localStorage.getItem('hr-lang');
    if (saved) setLang(saved);
    function filterTOC(q) {{
      const items = document.querySelectorAll('.toc-item');
      const noResult = document.getElementById('noResult');
      const search = q.toLowerCase().trim();
      let visible = 0;
      items.forEach(item => {{
        const text = item.textContent.toLowerCase();
        if (!search || text.includes(search)) {{
          item.classList.remove('hidden'); visible++;
        }} else {{
          item.classList.add('hidden');
        }}
      }});
      noResult.style.display = visible === 0 ? 'block' : 'none';
    }}
  </script>
</body>
</html>'''

    with open(toc_path, "w", encoding="utf-8") as f:
        f.write(toc_content)

    print(f"\n✅ อัปเดต TOC")
    print("\n" + "="*56)
    print("  เสร็จ! git add . && git commit -m 'v7 TH+EN' && git push")
    print("="*56 + "\n")

if __name__ == "__main__":
    main()
